import docx
import os
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

document = docx.Document() # 新建docx文档
style_song = document.styles.add_style('Song', WD_STYLE_TYPE.CHARACTER) # 设置Song字样式
style_song.font.name = '宋体'
document.styles['Song']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') # 将段落中的所有字体
style_song = document.styles.add_style('Kai', WD_STYLE_TYPE.CHARACTER)
style_song.font.name = '楷体'
document.styles['Kai']._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体') # 将段落中的所有字体
style_song = document.styles.add_style('Lishu', WD_STYLE_TYPE.CHARACTER)
style_song.font.name = '隶书'
document.styles['Lishu']._element.rPr.rFonts.set(qn('w:eastAsia'), u'隶书') # 将段落中的所有字体
paragraph1 = document.add_paragraph() # 添加段落
run = paragraph1.add_run(u'aBCDefg这是中文', style='Song') # 设置宋体样式
font = run.font #设置字体font.name = 'Cambira' # 设置西文字体
paragraph1.add_run(u'aBCDefg这是中文', style='Kai').font.name = 'Cambira'
paragraph1.add_run(u'aBCDefg这是中文', style='Lishu').font.name = 'Cambira'