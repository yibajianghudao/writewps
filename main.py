import docx
import os
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def read_save_document(file_path, output_file_path):
    doc = docx.Document(file_path)
    # 读取文件内所有段落
    paragraphs = [p.text for p in doc.paragraphs]
    # 保存要写入的段落
    str1 = ""
    doc = docx.Document()
    for paragraph in paragraphs:
        # print(f"p:{paragraph}")
        str1 += paragraph + '，'
    # print(f"{str1}")

    # 设置中文字体
    style_song = doc.styles.add_style('Song', WD_STYLE_TYPE.CHARACTER)
    style_song.font.name = 'Times New Roman'
    style_song.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 写入并保存到输出文件
    paragraph = doc.add_paragraph()
    paragraph.add_run(str1, style='Song')
    doc.save(output_file_path)
    print(f"{output_file_path}已经保存")

# 读取全部文章
def process_files(input_dir):
    file_list = os.listdir(input_dir)
    for file in file_list:
        input_file = os.path.join(input_dir, file)
        print(input_file)
        output_file = input_file.replace("字幕.docx", "") + "分段.docx"
        read_save_document(input_file, output_file)
# 读取目录下全部文件示例
# input_directory = r"/home/JiangHuDao/Downloads/wenzhang"
# output_directory = r"/home/JiangHuDao/Downloads/wenzhang"
# process_files(input_directory)

# 读取指定文件示例
input_file = "./test.docx"
output_file = "./tested.docx"
read_save_document(input_file, output_file)


